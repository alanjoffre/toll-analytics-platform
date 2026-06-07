# -*- coding: utf-8 -*-
"""
check_docs_sync.py — VERIFICADOR DE DRIFT da documentação.

Compara o documento Word (e o README) com o ESTADO REAL do código e falha (exit 1)
se a documentação ficou para trás. Pega os dois tipos de drift que já aconteceram:
  (1) um ADR citado no código/README que NÃO está explicado no Word;
  (2) um arquivo do projeto (.sql/.yml/.csv) que NÃO é mencionado no Word.

Uso:
    .venv/bin/python docs/gerar_documentacao.py     # 1) regenera o Word
    .venv/bin/python docs/check_docs_sync.py        # 2) confere se está em dia

Roda no CI (ver .github/workflows/dbt_ci.yml) e via `make check-docs`.
"""

import os
import re
import sys
import glob
from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX = os.path.join(ROOT, "docs", "Documentacao_dbt_toll_analytics.docx")

# Arquivos que NÃO precisam ser citados no Word (a própria tooling de docs).
IGNORE_BASENAMES = {
    "gerar_documentacao.py",
    "check_docs_sync.py",
    "Documentacao_dbt_toll_analytics.docx",
}


def docx_text(path):
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def find_adrs_in_code():
    """Todos os ADR-N citados no README e no código-fonte."""
    adrs = set()
    patterns = [
        "README.md",
        "PLANO_DO_PROJETO.md",
        "dbt_project.yml",
        "models/**/*.sql",
        "models/**/*.yml",
        "snapshots/*.sql",
        "macros/*.sql",
        "seeds/*.yml",
    ]
    for pat in patterns:
        for f in glob.glob(os.path.join(ROOT, pat), recursive=True):
            try:
                with open(f, encoding="utf-8") as fh:
                    adrs.update(re.findall(r"ADR-\d+", fh.read()))
            except (OSError, UnicodeDecodeError):
                pass
    return adrs


def find_project_files():
    """Arquivos-fonte do projeto cujo nome deve aparecer no Word."""
    files = set()
    for pat in [
        "seeds/*.csv",
        "seeds/*.yml",
        "models/**/*.sql",
        "models/**/*.yml",
        "snapshots/*.sql",
        "macros/*.sql",
    ]:
        for f in glob.glob(os.path.join(ROOT, pat), recursive=True):
            base = os.path.basename(f)
            if base not in IGNORE_BASENAMES:
                files.add(base)
    return files


def main():
    if not os.path.exists(DOCX):
        print(
            f"ERRO: Word não encontrado em {DOCX}. Rode: python docs/gerar_documentacao.py"
        )
        return 1

    txt = docx_text(DOCX)
    problems = []

    # (1) ADRs
    adrs = find_adrs_in_code()
    missing_adrs = sorted(a for a in adrs if a not in txt)
    if missing_adrs:
        problems.append(
            "ADRs citados no código/README mas AUSENTES no Word: "
            + ", ".join(missing_adrs)
        )

    # (2) Arquivos
    files = find_project_files()
    missing_files = sorted(
        f for f in files if f not in txt and f.rsplit(".", 1)[0] not in txt
    )
    if missing_files:
        problems.append(
            "Arquivos do projeto AUSENTES no Word: " + ", ".join(missing_files)
        )

    if problems:
        print("DRIFT DETECTADO — a documentação está atrás do código:\n")
        for p in problems:
            print("  - " + p)
        print(
            "\nConserto: atualize docs/gerar_documentacao.py e rode "
            "`python docs/gerar_documentacao.py`."
        )
        return 1

    print(
        f"OK — documentação em dia: {len(adrs)} ADRs e {len(files)} arquivos "
        "todos cobertos no Word."
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
