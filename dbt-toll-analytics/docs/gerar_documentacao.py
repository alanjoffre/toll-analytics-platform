# -*- coding: utf-8 -*-
"""
Gera a documentação didática completa do projeto dbt-toll-analytics em .docx.
Rodar:  .venv/bin/python docs/gerar_documentacao.py
Saída:  docs/Documentacao_dbt_toll_analytics.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

AZUL = RGBColor(0x1F, 0x4E, 0x79)
CINZA = RGBColor(0x59, 0x59, 0x59)
VERDE = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()

# ---- estilos base -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = AZUL
    return p

def para(text, italic=False, bold=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # fundo cinza claro
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "1F4E79")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()

# =============================================================================
# CAPA
# =============================================================================
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("dbt-toll-analytics"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = AZUL
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documentação Didática Completa do Projeto"); r.font.size = Pt(16); r.font.color.rgb = CINZA
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Engenharia de Dados com dbt + DuckDB — Analytics e Auditoria de Vale-Pedágio")
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = CINZA
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("O QUE foi feito · COMO foi feito · ONDE · POR QUÊ\nExplicação arquivo por arquivo + falas prontas para o LinkedIn")
r.font.size = Pt(11); r.font.color.rgb = AZUL
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documento gerado automaticamente a partir do projeto real (validado com dbt build verde).")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = CINZA
page_break()

# =============================================================================
# SUMÁRIO (campo TOC automático — atualiza no Word com F9 / botão direito)
# =============================================================================
h("Sumário", 1)
para("Para atualizar os números de página: clique com o botão direito no sumário "
     "abaixo e escolha “Atualizar campo” (ou selecione tudo e tecle F9).",
     italic=True, size=9, color=CINZA)
par = doc.add_paragraph()
run = par.add_run()
fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
t2 = OxmlElement('w:t'); t2.text = "  [O sumário aparece aqui depois de abrir no Word e atualizar o campo — botão direito > Atualizar campo]"
fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar); run._r.append(instrText); run._r.append(fldChar2); run._r.append(t2); run._r.append(fldChar3)
page_break()

# =============================================================================
# 1. VISÃO GERAL
# =============================================================================
h("1. Visão geral: o que é este projeto", 1)
para("Este projeto é um “produto de dados” completo de auditoria e analytics de "
     "vale-pedágio. Em uma frase: ele pega transações de pedágio cruas (sujas, com erros "
     "de propósito), limpa e organiza esses dados em camadas, e no final entrega duas "
     "coisas: (a) tabelas analíticas confiáveis para relatórios e (b) um modelo de "
     "auditoria que aponta automaticamente quais transações são suspeitas.")
para("Tudo roda 100% no seu computador, de graça, usando um banco de dados chamado "
     "DuckDB, e tudo é construído com a ferramenta dbt. Os dados são sintéticos "
     "(inventados) — nenhum dado real de cliente é usado, o que torna o projeto seguro "
     "para colocar no portfólio público (GitHub, LinkedIn).")

h("1.1. Quais problemas o modelo de auditoria detecta", 2)
table(
    ["Anomalia (flag)", "O que significa", "Exemplo nos dados"],
    [
        ["TARIFA_DIVERGENTE", "Valor cobrado diferente do esperado para aquela praça/veículo/data", "T0007, T0022, T0023"],
        ["VALOR_INVALIDO", "Cobrança com valor nulo ou zero", "T0010 (zero), T0026 (nulo)"],
        ["COBRANCA_EM_FALHA", "Transação marcada como FAILED/REVERSED mas com valor cobrado", "T0009 (FAILED), T0018 (REVERSED)"],
        ["POSSIVEL_DUPLICIDADE", "Mesmo veículo na mesma praça em poucos minutos (cobrança repetida)", "T0015 + T0016 (3 min)"],
    ],
    widths=[1.9, 3.1, 1.8],
)
para("Resultado validado: o pipeline encontra exatamente 9 transações suspeitas "
     "(3 tarifa divergente, 2 valor inválido, 2 possível duplicidade, 2 cobrança em falha).",
     bold=True, color=VERDE)

# =============================================================================
# 2. POR QUE DBT
# =============================================================================
h("2. Por que usamos dbt (e o que é)", 1)
para("dbt (data build tool) é a ferramenta padrão de mercado para a camada de "
     "TRANSFORMAÇÃO de dados — o “T” do ELT (Extract, Load, Transform). A ideia "
     "central: em vez de escrever scripts SQL soltos e bagunçados, você escreve SQL "
     "modular que roda dentro do próprio banco de dados, aplicando as boas práticas de "
     "engenharia de software ao trabalho com dados.")
h("2.1. O que o dbt traz de valor", 2)
bullet("você versiona tudo no Git, como código de verdade.", "Versionamento: ")
bullet("cada transformação é um arquivo .sql pequeno que referencia outro — o dbt "
       "monta sozinho a ordem de execução (o DAG).", "Modularidade: ")
bullet("você declara regras (“essa coluna não pode ser nula”, “essa chave é "
       "única”) e o dbt valida automaticamente.", "Testes de dados: ")
bullet("o dbt gera um site navegável com a descrição de cada tabela e o lineage "
       "(o mapa de dependências de ponta a ponta).", "Documentação + lineage: ")
bullet("o mesmo SQL que roda localmente no DuckDB roda no Databricks, "
       "Snowflake, BigQuery, etc. — basta trocar a conexão.", "Portabilidade: ")
para("Em resumo: dbt transforma “escrever SQL” em “fazer engenharia de dados”. "
     "É exatamente o que separa um trabalho amador de um pipeline de produção.")

h("2.2. Por que DuckDB (e a ponte para o Databricks)", 2)
para("DuckDB é um banco analítico que roda em um único arquivo local, sem servidor, "
     "sem nuvem e sem custo. É perfeito para aprender e demonstrar, porque o fluxo dbt é "
     "idêntico ao de um data lakehouse profissional. Quando quiser levar para produção no "
     "Databricks, troca-se apenas o arquivo de conexão (profiles.yml) e o adapter — os "
     "modelos SQL permanecem os mesmos. É isso que torna o aprendizado local 100% "
     "transferível para o trabalho real.")

# =============================================================================
# 3. ONDE FOI FEITO (ferramentas)
# =============================================================================
h("3. Onde e em quais programas foi feito", 1)
table(
    ["Programa / Ferramenta", "Para que serviu"],
    [
        ["VS Code", "Editor de código onde todos os arquivos foram escritos."],
        ["Terminal (zsh, macOS)", "Onde rodamos os comandos dbt e a validação."],
        ["Python 3.12 + venv", "Ambiente isolado para instalar o dbt e dependências."],
        ["dbt-core 1.11 + dbt-duckdb", "O motor que executa as transformações."],
        ["DuckDB", "O banco de dados local (arquivo toll_analytics.duckdb)."],
        ["Git / GitHub", "Versionamento e CI (GitHub Actions roda o pipeline em cada PR)."],
        ["SQLFluff", "Linter que verifica o estilo/qualidade do SQL."],
        ["Elementary (dbt package)", "Observabilidade de dados: detecção de anomalia + tabelas de monitoramento."],
        ["MetricFlow (mf)", "Runtime do Semantic Layer para consultar métricas."],
        ["python-docx", "Biblioteca que gerou este próprio documento Word."],
    ],
    widths=[2.4, 4.4],
)

# =============================================================================
# 4. ARQUITETURA MEDALLION
# =============================================================================
h("4. A arquitetura em camadas (Medallion)", 1)
para("O projeto organiza os dados em três camadas de qualidade crescente — o padrão "
     "“Medallion” (medalha): Bronze, Silver e Gold. Cada camada tem uma "
     "responsabilidade única, o que torna o pipeline fácil de entender, testar e manter.")
table(
    ["Camada", "Pasta no projeto", "Responsabilidade"],
    [
        ["BRONZE (cru)", "seeds/", "Dados crus, exatamente como chegaram — inclusive os erros propositais."],
        ["SILVER (limpo)", "models/staging/", "Tipagem, limpeza, deduplicação e mascaramento de PII. 1 modelo por fonte. Sem juntar fontes."],
        ["GOLD (consumo)", "models/intermediate/ + models/marts/", "Enriquecimento (joins + regra de negócio) e o modelo dimensional + auditoria."],
    ],
    widths=[1.5, 2.1, 3.2],
)
para("Fluxo visual (da esquerda para a direita):", bold=True)
code("seeds (bronze)  ->  staging (silver)  ->  intermediate  ->  marts (gold)\n"
     "raw_*.csv            stg_*.sql            int_*.sql         dim_*/fct_*/agg_*/audit_*")
para("A regra de ouro: a camada silver só “arruma a casa” (não junta fontes); "
     "quem junta e aplica regra de negócio é a camada gold. Isso mantém cada peça simples "
     "e reaproveitável.")

# =============================================================================
# 5. PASSO A PASSO DA CONSTRUÇÃO
# =============================================================================
page_break()
h("5. Passo a passo: como o projeto foi construído", 1)
para("Construímos respeitando o DAG (de baixo para cima): primeiro a base, depois o que "
     "depende dela, validando a cada bloco. Esta foi a ordem exata seguida:")
passos = [
    ("Passo 1 — Configuração", "Ajustamos packages.yml (bibliotecas dbt_utils e dbt_expectations) e requirements.txt (dbt-duckdb e sqlfluff). É a fundação: define quais ferramentas e bibliotecas o projeto usa."),
    ("Passo 2 — Seeds (bronze)", "Tiramos a tarifa fixa da tabela de praças e criamos raw_fare_schedule.csv, a tabela de tarifa COM VIGÊNCIA (com a mudança de preço da praça P003). Isso habilita a técnica point-in-time."),
    ("Passo 3 — Staging (silver)", "Ajustamos stg_toll_plazas, criamos stg_fare_schedule e o principal: stg_toll_transactions (deduplicação + tipagem + chave técnica). Documentamos e testamos tudo em _staging.yml."),
    ("Passo 4 — Snapshot", "Criamos snap_toll_plazas para demonstrar a captura de histórico (SCD2) nativa do dbt."),
    ("Passo 5 — Intermediate", "Criamos int_transactions_enriched: o coração do projeto, onde acontece o join point-in-time da tarifa e a detecção de duplicidade na janela de tempo."),
    ("Passo 6 — Macros", "Criamos cents_to_brl (converte centavos para reais), audit_flag (a regra de auditoria) e log_run_results (observabilidade)."),
    ("Passo 7 — Marts (gold)", "Criamos as dimensões (dim_date, dim_plaza, dim_vehicle), o fato incremental (fct_toll_transactions), a agregação de receita e o produto final audit_suspect_transactions. Aplicamos contratos de dados e unit tests."),
    ("Passo 8 — Testes", "Criamos um teste genérico customizado (not_charged_when_failed) e um teste singular de unicidade no fato."),
    ("Passo 9 e 10 — Exposures e CI", "Declaramos o consumidor downstream (relatório aos gestores) no lineage, e configuramos o CI (GitHub Actions) + o linter SQLFluff."),
    ("Passo 11 — README com ADRs", "Reescrevemos o README documentando as decisões de arquitetura (ADRs)."),
    ("Passo 12 — Validação", "Criamos o ambiente Python, instalamos as dependências e rodamos dbt build — resultado verde (PASS=92, WARN=1 intencional, ERROR=0)."),
    ("Passo 13 — Hardening sênior", "Após um code review crítico, corrigimos 3 bugs reais (fan-out no join, late-arriving data, dim_date fixa) e adicionamos diferenciais (Semantic Layer, dev/prod). Ver seções 9.1 e 9.2."),
]
for titulo, desc in passos:
    p = doc.add_paragraph()
    r = p.add_run(titulo + ": "); r.bold = True; r.font.color.rgb = AZUL
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(6)

# =============================================================================
# 6. EXPLICAÇÃO ARQUIVO POR ARQUIVO
# =============================================================================
page_break()
h("6. Explicação de cada arquivo (o que é, para que serve, a lógica)", 1)
para("Esta é a parte mais detalhada: cada arquivo do projeto, agrupado por pasta, com a "
     "lógica por trás dele.")

def arquivo(nome, o_que, porque, logica=None):
    p = doc.add_paragraph()
    r = p.add_run(nome); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = AZUL
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    q = doc.add_paragraph(); q.add_run("O que é / para que serve: ").bold = True; q.add_run(o_que)
    q.paragraph_format.space_after = Pt(2)
    w = doc.add_paragraph(); w.add_run("Por que existe: ").bold = True; w.add_run(porque)
    w.paragraph_format.space_after = Pt(2)
    if logica:
        l = doc.add_paragraph(); l.add_run("Lógica interna: ").bold = True; l.add_run(logica)
        l.paragraph_format.space_after = Pt(4)

h("6.1. Raiz do projeto (configuração)", 2)
arquivo("dbt_project.yml",
        "O arquivo central de configuração do dbt. Define o nome do projeto, onde ficam as pastas, e as configurações globais.",
        "Todo projeto dbt precisa dele — é o ponto de entrada que o dbt lê primeiro.",
        "Define: vars (variáveis de negócio, como a janela de 300s para duplicidade), o gancho on-run-end (chama log_run_results ao fim de cada execução), store_failures (materializa linhas que falham em testes), as materializações por camada (staging=view, marts=table) e a tipagem das datas do seed de tarifa.")
arquivo("profiles.yml",
        "O arquivo de conexão com o banco, com DOIS ambientes: dev (default) e prod (banco/schema separados).",
        "Separa a conexão do código e demonstra separação de ambientes (dbt build --target prod). É exatamente este arquivo que se troca para apontar ao Databricks, sem mexer nos modelos.")
arquivo("packages.yml",
        "Lista as bibliotecas externas de dbt usadas: dbt_utils e dbt_expectations.",
        "Reaproveitar código testado pela comunidade (geração de chaves, série de datas, testes avançados) em vez de reinventar.")
arquivo("requirements.txt",
        "Lista os pacotes Python necessários: dbt-duckdb, duckdb e sqlfluff.",
        "Garante que qualquer pessoa (ou o CI) instale exatamente as mesmas ferramentas e reproduza o projeto.")
arquivo(".gitignore",
        "Diz ao Git o que NÃO versionar (o banco .duckdb, a pasta target/, o .venv/).",
        "Evita subir arquivos gerados/binários para o repositório — boa higiene de Git.")
arquivo(".sqlfluff",
        "Configuração do linter de SQL (dialeto duckdb, regras de estilo, indentação).",
        "Padroniza o estilo do SQL no time e no CI — qualidade automatizada.")

h("6.2. Ingestão (dlt) — camada BRONZE (landing)", 2)
para("O bronze NÃO é mais seed (ADR-28): o projeto ingestion-toll-analytics usa o "
     "dlt para ler estes arquivos de landing (CSV) e carregá-los no schema 'landing' "
     "do DuckDB; o staging consome via source('toll_raw', ...). Os arquivos abaixo "
     "vivem em ingestion-toll-analytics/data/.")
arquivo("raw_toll_plazas.csv",
        "As praças de pedágio (id, nome, rodovia, UF).",
        "É a dimensão de local. A tarifa NÃO fica mais aqui — foi movida para o schedule (ver ADR-1), porque preço tem história.")
arquivo("raw_fare_schedule.csv",
        "A tabela de tarifa COM VIGÊNCIA: cada linha tem plaza_id, fare_cents, valid_from, valid_to.",
        "É a fonte da verdade para o join point-in-time. Contém a mudança proposital: a praça P003 passa de 990 para 1050 centavos em 2026-05-03.",
        "Modelar tarifa com vigência é o estilo SCD2 (Slowly Changing Dimension tipo 2) — guarda o histórico de preço.")
arquivo("raw_vehicles.csv",
        "Os veículos (id, placa, categoria, conta).",
        "A placa é PII (dado pessoal) e será mascarada na camada silver, por LGPD.")
arquivo("raw_vehicle_categories.csv",
        "As categorias de veículo e o multiplicador de tarifa (carro 1.0x, carreta 4.5x, etc.).",
        "O valor esperado de uma transação depende da categoria: tarifa da praça x multiplicador.")
arquivo("raw_toll_transactions.csv",
        "As passagens de pedágio (id, veículo, praça, timestamp, valor, forma de pagamento, status).",
        "É a tabela-fato crua. Contém TODOS os erros propositais (duplicata exata, valor nulo/zero, tarifa divergente, cobrança em falha) que o pipeline vai limpar e auditar.")
arquivo("toll_ingestion.py (dlt) + _sources.yml",
        "O pipeline dlt (projeto ingestion-toll-analytics) carrega os arquivos de landing no schema 'landing' (merge/replace, metadados de carga, '' -> NULL); o _sources.yml declara a source toll_raw (schema landing) + freshness + testes.",
        "É o EL antes do T (ADR-28): o staging troca ref(seed) por source('toll_raw', ...). Reprodutível offline (o dlt lê CSVs commitados).")

h("6.3. models/staging/ — camada SILVER (limpo)", 2)
arquivo("stg_toll_plazas.sql",
        "Limpa e tipa as praças (cast, trim, upper).",
        "Padroniza os dados de praça. Removemos a coluna de tarifa daqui (ela vive no schedule).")
arquivo("stg_fare_schedule.sql (+ teste mutually_exclusive_ranges)",
        "Tipa a tabela de tarifa com vigência e — crucial — tem um teste que impede vigências sobrepostas por praça.",
        "Prepara o schedule para o join temporal. O teste dbt_utils.mutually_exclusive_ranges trava o fan-out: se duas vigências da mesma praça se sobrepusessem, o join duplicaria a transação silenciosamente. Ver ADR-9.")
arquivo("stg_vehicles.sql",
        "Limpa veículos e MASCARA a placa (mantém só os 3 primeiros caracteres + ****).",
        "Proteção de PII (LGPD) feita o mais cedo possível na pipeline, reduzindo a exposição do dado sensível.")
arquivo("stg_vehicle_categories.sql",
        "Limpa e tipa as categorias e o multiplicador.",
        "Padroniza os dados de categoria para o cálculo da tarifa esperada.")
arquivo("stg_toll_transactions.sql",
        "O modelo mais importante da silver: deduplica, tipa, deriva a data do evento e cria uma chave técnica (surrogate key).",
        "É aqui que a duplicata exata (T0020 repetida) é colapsada em uma única linha. Valores nulos/zero são MANTIDOS de propósito (regra: flag, não delete).",
        "Usa row_number() particionado por transaction_id para manter só 1 linha por transação. Gera transaction_sk com dbt_utils.generate_surrogate_key (hash do id).")
arquivo("_staging.yml",
        "Documenta e testa toda a camada silver (unique, not_null, relationships, accepted_values, accepted_range).",
        "Garante que a limpeza funcionou — por exemplo, que transaction_id é único DEPOIS do dedup.")
arquivo("_sources.yml",
        "Declara os SOURCES (contrato de ingestão da camada bronze) e o controle de FRESHNESS (atraso de ingestão) na tabela de transações.",
        "Mostra que entendemos contrato de ingestão, não só transformação. Em dev os sources apontam para as tabelas dos seeds; em prod, viriam de ingestão real (Auto Loader/COPY INTO) e o staging trocaria ref() por source(). Ver ADR-13.",
        "Roda com: dbt source freshness. Contra o seed estático reporta WARN (esperado); com ingestão real fica verde.")

h("6.4. models/intermediate/ — enriquecimento (GOLD)", 2)
arquivo("int_transactions_enriched.sql",
        "O CORAÇÃO do projeto. Junta as transações com veículo, categoria, praça e — o ponto alto — a tarifa vigente NA DATA do evento (point-in-time).",
        "Calcula a tarifa esperada (tarifa vigente x multiplicador), a diferença, e marca duplicidade na janela de tempo. É a base do fato e da auditoria.",
        "(1) Detecta duplicidade com lag()/lead() sobre veículo+praça, marcando AMBAS as passagens do par. (2) Faz o join point-in-time: event_date BETWEEN valid_from AND valid_to, pegando a tarifa certa para a data. (3) Usa LEFT JOIN para ser conservador (sem tarifa = não acusa divergência).")
arquivo("_intermediate.yml",
        "Documenta e testa o modelo enriquecido.",
        "Garante unicidade da transação e documenta os campos calculados (tarifa esperada, diferença, duplicidade).")
arquivo("int_duplicate_flags.sql (EPHEMERAL)",
        "Detecção de duplicidade na janela, materializada como EPHEMERAL: não vira objeto no banco — o dbt inlina como CTE em quem der ref().",
        "Separa 'achar duplicata' de 'enriquecer', deixando o int_transactions_enriched mais limpo. Uso clássico de ephemeral (passo lógico reutilizável e barato). Ver ADR-20.")

h("6.5. models/marts/ — camada GOLD (consumo)", 2)
arquivo("_groups.yml (groups + access)",
        "Define os GROUPS (staging/intermediate/marts) com dono e o ACCESS de cada camada: interno 'protected', marts 'public'.",
        "Governança de modelos: deixa explícito o que é consumível (marts public) e o que é interno do package (protected). Ver ADR-18.")
arquivo("dim_date.sql",
        "Dimensão de calendário gerada com dbt_utils.date_spine. A faixa NÃO é fixa: é derivada do min/max real das transações (ver ADR-10).",
        "Permite analisar por dia, mês, dia da semana, fim de semana. Derivar a faixa dos dados evita que uma transação fora de 2026 quebre o teste de relationships. Também serve de time spine para o Semantic Layer.")
arquivo("dim_plaza.sql",
        "Dimensão de praça, com surrogate key.",
        "Descreve a praça para os relatórios. A surrogate key desacopla a chave técnica da chave de negócio.")
arquivo("dim_vehicle.sql",
        "Dimensão de veículo (placa mascarada, categoria, multiplicador, conta).",
        "Descreve o veículo para análise, já com a PII protegida.")
arquivo("fct_toll_transactions.sql",
        "A tabela-fato (uma linha por transação), materializada como INCREMENTAL.",
        "Em escala, não reprocessa o histórico — só carrega o recente. unique_key garante idempotência. Carrega também a coluna audit_flag (mesma macro do produto) para habilitar métricas no Semantic Layer.",
        "Usa is_incremental() com janela de LOOKBACK (event_date >= max - N dias) para capturar transações atrasadas (late-arriving) sem descartá-las — o merge por unique_key deduplica. Ver ADR-8.")
arquivo("agg_daily_revenue_by_plaza.sql",
        "Agregação: receita diária por praça (somando só cobranças legítimas).",
        "Entrega uma visão de negócio pronta. Transações em falha ou inválidas NÃO entram na receita.")
arquivo("audit_suspect_transactions.sql",
        "O PRODUTO DE DADOS final: aplica a regra de auditoria e entrega só as transações suspeitas.",
        "É a entrega de valor do projeto. Usa a macro audit_flag e filtra onde a flag é diferente de OK.")
arquivo("py_plaza_audit_stats.py (PYTHON MODEL)",
        "Python model (dbt-duckdb): estatística de auditoria por praça (total, suspeitas, taxa, vazamento e z-score da taxa entre praças), em pandas.",
        "Mostra que o dbt orquestra/testa/versiona Python igual a SQL. Estatística (z-score entre praças) é natural em pandas. Ver ADR-21.")
arquivo("rpt_plaza_revenue_v1.sql / rpt_plaza_revenue_v2.sql (VERSIONED)",
        "Model VERSIONADO de report: v1 (deprecada em 2026-12-31) e v2 (latest, adiciona ticket médio).",
        "Quebra de contrato via VERSÃO, não in-place: consumidores migram no seu ritmo. ref('rpt_plaza_revenue') resolve para a latest_version (v2). Ver ADR-22.")
arquivo("_marts.yml",
        "Documenta e testa os marts, e aplica os CONTRATOS DE DADOS (trava colunas + tipos).",
        "O contrato faz o build QUEBRAR se alguém mudar o schema sem querer — proteção forte contra regressão. Também aplica o teste de severidade warn.")
arquivo("_unit_tests.yml",
        "Unit tests (dbt 1.8+): testam a LÓGICA da auditoria com inputs inventados e saída esperada.",
        "Diferente dos testes de dados (que olham os dados reais), o unit test prova que a REGRA funciona, de forma determinística. Pouquíssimo portfólio tem isso.",
        "Dois testes: um cobre cada tipo de anomalia; o outro cobre a PRIORIDADE entre regras (ex.: valor inválido vence cobrança em falha).")
arquivo("_semantic_models.yml",
        "O Semantic Layer (MetricFlow): define a camada semântica sobre o fato e as MÉTRICAS governadas (revenue, transactions, suspect_rate, revenue_leakage_brl).",
        "Em vez de uma tabela agregada por recorte, define a métrica UMA vez e o MetricFlow gera o SQL sob demanda para qualquer dimensão. Garante consistência da métrica e é a feature dbt mais 'sênior 2025/2026'. Ver ADR-11.",
        "Consulta: mf query --metrics revenue,suspect_rate --group-by metric_time__day. Validado de verdade.")
arquivo("exposures.yml",
        "Declara o consumidor downstream: o 'Relatório de auditoria aos gestores'.",
        "Deixa explícito no lineage QUEM é afetado se um modelo upstream quebrar — gestão de impacto.")

h("6.6. snapshots/, macros/ e tests/", 2)
arquivo("snapshots/snap_toll_plazas.sql",
        "Snapshot SCD2 nativo do dbt: captura mudanças de atributo da praça ao longo do tempo.",
        "Demonstra a técnica de CDC/SCD2 do dbt. Nota: para a tarifa usamos o seed schedule (história determinística já na 1a execução) — ver ADR-2.")
arquivo("macros/cents_to_brl.sql",
        "Macro que converte centavos para reais (valor numérico, 2 casas).",
        "Centraliza a conversão para não espalhar '/100' mágico pelo código.")
arquivo("macros/audit_flag.sql",
        "A REGRA DE NEGÓCIO da auditoria, isolada numa macro reutilizável e testável.",
        "Classifica cada transação em uma flag por ordem de prioridade. Isolar na macro permite testá-la com unit tests.",
        "Prioridade: VALOR_INVALIDO > COBRANCA_EM_FALHA > TARIFA_DIVERGENTE > POSSIVEL_DUPLICIDADE > OK.")
arquivo("macros/log_run_results.sql",
        "Observabilidade: ao fim de cada execução, grava metadados (status, tempo, linhas) na tabela _audit_runs.",
        "Permite monitorar o pipeline ao longo do tempo — prática de produção. Chamada pelo gancho on-run-end.")
arquivo("tests/generic/test_not_charged_when_failed.sql",
        "Teste genérico customizado e reutilizável: 'não deveria haver valor cobrado em transação FAILED/REVERSED'.",
        "Demonstra Data Quality como sistema (teste reusável), não verificação ad-hoc. Aplicado com severidade warn (ver ADR-6).")
arquivo("tests/assert_unique_transaction_in_fct.sql",
        "Teste singular: garante que o fato não tem transaction_id duplicado.",
        "Redundância proposital (cinto e suspensório) confirmando que dedup + incremental funcionaram.")

h("6.7. CI/CD e documentação", 2)
arquivo(".github/workflows/dbt_ci.yml",
        "CI de PR no GitHub Actions: lint (models+tests+snapshots+macros), dbt build --exclude tag:observability, source freshness e dbt docs generate.",
        "Qualidade automatizada = produção. Tem cache de pip e publica o lineage como artefato. Exclui a tag observability para o PR ficar determinístico (ADR-17). Roadmap: Slim CI + GitHub Pages.")
arquivo(".github/workflows/observability.yml",
        "Workflow AGENDADO (nightly) de observabilidade: roda a detecção de anomalia (dbt test --select tag:observability) com o autoupload do Elementary ligado.",
        "Anomalia precisa de HISTÓRICO acumulado de várias execuções — por isso roda agendado, não a cada PR. Separa monitoramento contínuo de gate de merge (ADR-17).")
arquivo("README.md",
        "A porta de entrada do repositório: o que o projeto faz, como rodar, e os ADRs (decisões de arquitetura).",
        "Documentar o julgamento por trás das decisões é o que mais sinaliza senioridade.")
arquivo("PLANO_DO_PROJETO.md",
        "O 'cérebro' do projeto: explica o quê, o porquê e o como de cada técnica.",
        "Permite retomar o projeto em qualquer sessão futura e serve de roteiro de estudo.")

# =============================================================================
# 7. A TÉCNICA-ESTRELA
# =============================================================================
page_break()
h("7. A técnica-estrela: tarifa point-in-time", 1)
para("Esta é a parte que mais impressiona em entrevista, então vale entender bem.")
para("O problema: o preço do pedágio muda no tempo. Se eu comparar uma transação ANTIGA "
     "contra a tarifa de HOJE, vou gerar um falso positivo de 'tarifa divergente' — "
     "acusar como errado algo que estava certo na época.")
para("A solução: juntar a transação à tabela de tarifa pela DATA DO EVENTO "
     "(event_date BETWEEN valid_from AND valid_to), pegando a tarifa que estava vigente "
     "naquele dia específico.")
para("Exemplo real nos dados (praça P003, carreta categoria 9, multiplicador 4.5x; a "
     "tarifa muda de 990 para 1050 em 2026-05-03):", bold=True)
table(
    ["Transação", "Data", "Valor cobrado", "Esperado (point-in-time)", "Resultado"],
    [
        ["T0006", "01/05", "4455", "990 x 4.5 = 4455", "OK (não acusa)"],
        ["T0023", "03/05", "4455", "1050 x 4.5 = 4725", "TARIFA_DIVERGENTE"],
    ],
    widths=[1.2, 0.9, 1.3, 2.0, 1.4],
)
para("Repare: a MESMA transação (valor 4455) é correta em 01/05 e divergente em 03/05, "
     "porque a tarifa vigente mudou. Um modelo júnior (que junta na tarifa atual) marcaria "
     "T0006 erroneamente como divergente. O join point-in-time é o que evita isso.",
     italic=True)

# =============================================================================
# 8. RESULTADOS VALIDADOS
# =============================================================================
h("8. Resultados validados (dbt build verde)", 1)
para("Rodamos o pipeline completo de verdade. Saída do dbt build do CI de PR "
     "(que exclui a tag observability — ela roda num job agendado, ADR-17):")
code("dbt build --exclude tag:observability\nDone. PASS=124  WARN=1  ERROR=0  SKIP=0  NO-OP=1  TOTAL=126")
bullet("124 verificações passaram (testes de dados + unit tests + ranges + audit_flag + modelos do Elementary). A detecção de anomalia roda separada: dbt test --select tag:observability.")
bullet("1 WARN é INTENCIONAL: o teste not_charged_when_failed pega as cobranças em "
       "FAILED/REVERSED que existem de propósito — a auditoria as sinaliza, então é "
       "alerta monitorável (warn), não falha de build.")
bullet("0 erros.")
bullet("Idempotência confirmada: rodar duas vezes mantém o fato com 26 linhas (o "
       "incremental não duplica).")
para("Distribuição final da auditoria: 3 TARIFA_DIVERGENTE, 2 VALOR_INVALIDO, "
     "2 POSSIVEL_DUPLICIDADE, 2 COBRANCA_EM_FALHA = 9 transações suspeitas.",
     bold=True, color=VERDE)

# =============================================================================
# 9. ADRs
# =============================================================================
h("9. Decisões de arquitetura (ADRs) — o porquê de cada escolha", 1)
para("ADR = Architecture Decision Record. Registrar o julgamento por trás das decisões é "
     "o que mais sinaliza senioridade.")
adrs = [
    ("ADR-1 — Tarifa no schedule, não na praça", "Preço tem história; modelar com vigência permite o join point-in-time e evita falso positivo."),
    ("ADR-2 — Histórico de tarifa via seed, não snapshot", "Snapshot só constrói história ao longo de várias execuções; para um portfólio reprodutível, precisamos da história já na 1a execução."),
    ("ADR-3 — Valor inválido: flag, não delete", "Em auditoria, a regra é sinalizar e manter (rastreabilidade), nunca apagar o dado."),
    ("ADR-4 — Fato incremental com unique_key", "Em escala não reprocessa o histórico; unique_key garante idempotência."),
    ("ADR-5 — Regra de auditoria isolada na macro", "Regra de negócio testável (unit tests) e reutilizável, desacoplada dos dados reais."),
    ("ADR-6 — Severidade warn para cobrança em falha", "São exatamente o que a auditoria existe para achar; quebrar o build seria contraditório. error fica para invariantes que nunca podem ocorrer."),
    ("ADR-7 — DuckDB local, portável para Databricks", "Loop de feedback local instantâneo e sem custo; os modelos SQL são os mesmos na nuvem."),
    ("ADR-8 — Incremental com janela de lookback", "Reprocessa os últimos N dias para capturar transações atrasadas (late-arriving); o unique_key deduplica. Evita o bug clássico de descartar dados que chegam fora de ordem."),
    ("ADR-9 — Guarda contra fan-out no join temporal", "Teste mutually_exclusive_ranges impede vigências sobrepostas por praça, que duplicariam transações silenciosamente."),
    ("ADR-10 — dim_date com faixa derivada dos dados", "Em vez de fixa em 2026, deriva do min/max real — robusto a datas fora do ano."),
    ("ADR-11 — Semantic Layer em vez de só agregações fixas", "Métricas definidas uma vez e consultáveis por qualquer dimensão; garante consistência da métrica."),
    ("ADR-12 — audit_flag materializada no fato", "Habilita métricas de suspeita no Semantic Layer sem duplicar regra (a lógica vive na macro — DRY)."),
    ("ADR-13 — sources + freshness sobre os seeds", "Declara o contrato de ingestão (SLA de atraso) sem perder o bronze reprodutível offline; em prod, seeds viram tabelas de ingestão e o staging passa a usar source()."),
    ("ADR-14 — Observabilidade de dados com Elementary", "Pacote elementary + teste de anomalia de volume; tabelas de monitoramento materializadas no warehouse. Salto de 'tenho metadados' para 'monitoro dados'. Ressalvas honestas no DuckDB: anomaly detection precisa de histórico (fica em warn); relatório HTML do edr é melhor em warehouses cloud."),
    ("ADR-15 — Dinheiro agregado em centavos inteiros", "Agregações monetárias somam amount_cents (inteiro, exato) e convertem para BRL só no final; nunca somam reais já arredondados. Evita acúmulo de erro de arredondamento/float — padrão em dado financeiro. Vale para o agg e para o Semantic Layer."),
    ("ADR-16 — Testes genéricos no formato arguments:", "Todos os testes passaram a aninhar parâmetros sob 'arguments:', eliminando os deprecation warnings do dbt 1.11+ (à prova do dbt 2.0). dbt parse roda sem nenhum aviso."),
    ("ADR-17 — Observabilidade como job agendado (fora do PR)", "O teste de anomalia ganhou a tag 'observability'; o CI de PR roda 'dbt build --exclude tag:observability' (determinístico, PASS=124) e um workflow agendado roda a detecção de anomalia. Motivo: anomalia precisa de histórico acumulado e o autoupload do Elementary no DuckDB dispara erro de commit. Monitoramento contínuo não é gate de merge."),
    ("ADR-18 — Groups + access (governança de modelos)", "Cada model pertence a um group (staging/intermediate/marts) com dono, e tem nível de access: o interno fica 'protected' (só o package referencia) e os marts ficam 'public' (camada consumível por BI/Semantic Layer/exposures). Documenta as fronteiras de consumo. 'private' não se aplica porque os refs cruzam grupos."),
    ("ADR-19 — Constraints no warehouse (PK/CHECK via contract)", "Com contract enforced, declaramos constraints (primary_key, not_null, check) que viram DDL REAL na CREATE TABLE — o banco passa a garantir a invariante, não só o teste dbt. PK no grão do fato e das dims; CHECK em dim_date (mês 1..12, dia da semana 0..6)."),
    ("ADR-20 — Model ephemeral (int_duplicate_flags)", "A detecção de duplicidade virou um model EPHEMERAL: não materializa objeto no banco — o dbt inlina como CTE em quem der ref(). Separa 'achar duplicata' de 'enriquecer'. Uso clássico de ephemeral: passo lógico reutilizável e barato que não precisa ser consultado isoladamente."),
    ("ADR-21 — Python model (py_plaza_audit_stats)", "Um model em PYTHON (dbt-duckdb) calcula estatística de auditoria por praça (taxa de suspeita e z-score entre praças) em pandas. O dbt orquestra/testa/versiona Python igual a SQL; usamos Python onde é natural (estatística), não por moda."),
    ("ADR-22 — Versioned model + deprecation_date", "rpt_plaza_revenue é VERSIONADO: a v2 (latest) adiciona ticket médio; a v1 fica deprecada até 2026-12-31. Quebra de contrato via VERSÃO, não in-place — os consumidores migram no seu ritmo. ref() resolve para a latest_version."),
    ("ADR-23 — dbt_project_evaluator como warn", "Adicionamos o pacote que audita o PRÓPRIO projeto contra best practices (naming, fanout, undocumented, modelos públicos sem contract...). As descobertas ficam como 'warn' — recomendações de melhoria contínua monitoráveis, não gate de merge."),
    ("ADR-24 — Materializações/grants só no target prod (Databricks)", "materialized_view, estratégia incremental microbatch e grants são features do warehouse de produção (Databricks/Delta), não do DuckDB single-node de dev. Em vez de falsear no dev, ficam documentadas/config no target prod. No DuckDB rodam as materializações que ele suporta (view/table/incremental/ephemeral)."),
    ("ADR-25 — Docs blocks + persist_docs", "Descrições reutilizáveis viram docs blocks ({% docs %} em models/docs.md, referenciados com doc()), e +persist_docs empurra as descrições para COMMENTs no banco. Documentação versionada, sem duplicação, que vive junto do dado. Validado: o COMMENT do audit_flag no DuckDB traz o docs block renderizado."),
    ("ADR-26 — Saved queries + exports (Semantic Layer)", "Uma saved query (revenue_daily) agrupa métricas + recorte + export no Semantic Layer — consulta governada e reaproveitável, em vez de cada dashboard copiar SQL. O export materializa a métrica numa tabela (dbt Cloud / mf export); em core a definição é versionada e validada (mf validate-configs)."),
    ("ADR-27 — Slim CI (state:modified+ --defer)", "Um workflow de PR (dbt_slim_ci.yml) constrói só o que mudou + downstream, deferindo o resto a um baseline (manifest da branch base). Reconstruir tudo a cada PR é caro; o Slim CI roda em segundos quando pouca coisa muda. Validado: dbt ls --select state:modified+ seleciona exatamente o model alterado. Em prod, o --defer aponta para o warehouse."),
    ("ADR-28 — Ingestão (EL) com dlt; bronze deixa de ser seed", "Realiza o ADR-13: um pipeline dlt (projeto ingestion-toll-analytics) lê os arquivos de landing (CSV) e carrega no schema 'landing' do DuckDB (merge/replace, metadados de carga, '' -> NULL). O staging passa a consumir via source('toll_raw', ...) em vez de seeds; os seeds foram removidos. Continua reprodutível offline (o dlt lê CSVs commitados, não uma API). O Airflow roda a ingestão antes do transform; o CI também (dlt -> dbt build)."),
]
for titulo, desc in adrs:
    p = doc.add_paragraph()
    r = p.add_run(titulo + ": "); r.bold = True; r.font.color.rgb = AZUL
    p.add_run(desc); p.paragraph_format.space_after = Pt(5)

# --- 9.1 Correções (hardening) ---
h("9.1. Rodada de correções: os 3 bugs reais que encontramos e consertamos", 2)
para("Depois de um code review crítico (o tipo de olhar que separa pleno de sênior), "
     "encontramos e corrigimos três bugs REAIS. Reconhecer e consertar isto vale mais "
     "que dez features novas.")
table(
    ["Bug", "Risco", "Conserto"],
    [
        ["Fan-out no join point-in-time", "Se duas vigências de tarifa se sobrepusessem, a transação duplicaria silenciosamente.", "Teste mutually_exclusive_ranges (ADR-9)."],
        ["Late-arriving data no incremental", "Filtro '> max(data)' descartava PARA SEMPRE transações que chegam atrasadas.", "Janela de lookback de N dias + merge (ADR-8)."],
        ["dim_date travada em 2026", "Qualquer transação fora de 2026 quebraria o teste de relationships.", "Faixa derivada do min/max real (ADR-10)."],
        ["Dinheiro somado em float", "Somar reais já arredondados (amount_brl) acumula erro de arredondamento.", "Somar centavos inteiros e converter no final (ADR-15)."],
    ],
    widths=[2.0, 3.0, 1.8],
)
para("Nota de processo: um segundo code review apontou alegações que JÁ estavam corrigidas "
     "(lookback, Semantic Layer, sources/freshness, Elementary, dev/prod) — verificamos "
     "contra o código antes de agir. Acertou em um bug real (dinheiro em float, ADR-15) e "
     "na oportunidade de endurecer o CI. Verificar cada alegação contra o código, em vez de "
     "aceitar ou descartar, é parte da postura sênior.", italic=True)

# --- 9.2 Semantic Layer + Limitações ---
h("9.2. Semantic Layer (MetricFlow) e o roadmap honesto", 2)
para("Adicionamos o Semantic Layer do dbt — a feature mais 'sênior 2025/2026'. Em vez "
     "de tabelas agregadas fixas, definimos MÉTRICAS governadas (uma definição central, "
     "consultável por qualquer dimensão): revenue, transactions, suspect_rate e "
     "revenue_leakage_brl (vazamento de receita por subcobrança). Validamos com mf query "
     "de verdade — exemplo: taxa de suspeita de 27% a 50% por dia.")
para("Limitações conhecidas e roadmap (documentar o próprio risco é o movimento que mais "
     "impressiona um recrutador técnico):", bold=True)
bullet("Late-arriving além da janela de N dias: mitigado e documentado (ADR-8).")
bullet("sources + freshness: IMPLEMENTADO (ADR-13) — contrato de ingestão declarado, com SLA de atraso; validado com dbt source freshness.")
bullet("Observabilidade avançada (Elementary): IMPLEMENTADO (ADR-14) — pacote + detecção de anomalia + tabelas de monitoramento no warehouse. Ressalva honesta: relatório HTML do edr é melhor em warehouses cloud que no DuckDB.")
bullet("Orquestração (Airflow + Cosmos): a camada de agendamento/observabilidade por cima fecharia o fim-a-fim. Próximo passo.")
bullet("Docs hospedados: publicar o dbt docs no GitHub Pages para lineage navegável linkável no LinkedIn.")

# =============================================================================
# 10. COMO RODAR
# =============================================================================
h("10. Como rodar o projeto na sua máquina", 1)
code("cd dbt-toll-analytics\n"
     "python3 -m venv .venv && source .venv/bin/activate\n"
     "pip install -r requirements.txt\n"
     "dbt deps  --profiles-dir .\n"
     "dbt build --profiles-dir .          # seed + run + snapshot + test\n"
     "dbt docs generate --profiles-dir . && dbt docs serve --profiles-dir .   # lineage")
para("Para ver o produto de auditoria:")
code("python -c \"import duckdb; c=duckdb.connect('toll_analytics.duckdb'); "
     "print(c.sql('select audit_flag, count(*) from main.audit_suspect_transactions group by 1'))\"")

# =============================================================================
# 11. GLOSSÁRIO
# =============================================================================
page_break()
h("11. Glossário rápido (para fixar o vocabulário)", 1)
glos = [
    ("ELT", "Extract, Load, Transform — primeiro carrega o dado cru, depois transforma dentro do banco. O dbt cuida do T."),
    ("Medallion", "Arquitetura em 3 camadas: bronze (cru), silver (limpo), gold (consumo)."),
    ("Model", "Um arquivo .sql com um SELECT; o dbt o transforma em view ou tabela."),
    ("ref()", "Função do dbt que referencia outro modelo e monta o DAG (ordem de execução)."),
    ("DAG / lineage", "O grafo de dependências entre os modelos, gerado automaticamente."),
    ("Materialização", "Como o modelo vira objeto: view, table, incremental, ephemeral."),
    ("Seed", "Um CSV versionado que o dbt carrega como tabela (nossa camada bronze)."),
    ("Snapshot", "Captura histórico de mudanças ao longo do tempo (SCD2 nativo)."),
    ("Surrogate key", "Chave técnica (hash) que desacopla a chave do negócio."),
    ("Point-in-time", "Join que usa a versão do dado vigente na data do evento."),
    ("Contract", "Contrato de dados que trava colunas e tipos; quebra o build se mudar."),
    ("Unit test", "Teste da LÓGICA do modelo com inputs mockados (dbt 1.8+)."),
    ("Exposure", "Declaração de um consumidor downstream (dashboard) no lineage."),
    ("PII", "Dado pessoal identificável (ex.: placa) — mascarado por LGPD."),
    ("Idempotência", "Rodar de novo produz o mesmo resultado, sem duplicar."),
    ("CI", "Integração Contínua — roda o pipeline automaticamente em cada mudança."),
]
table(["Termo", "Significado"], glos, widths=[1.6, 5.2])

# =============================================================================
# 12. LINKEDIN
# =============================================================================
h("12. Como traduzir isto para o seu LinkedIn", 1)
para("Abaixo, falas prontas (técnica -> o que dizer). Use no resumo, na descrição da "
     "experiência e em entrevistas. São honestas: descrevem o que o projeto realmente faz.")
table(
    ["Técnica no projeto", "Fala pronta para o LinkedIn / entrevista"],
    [
        ["Point-in-time (SCD2 via schedule)", "“Modelei tarifa com vigência e join point-in-time para evitar falso positivo em transações históricas.”"],
        ["Model incremental", "“Materializei o fato como incremental com unique_key para não reprocessar histórico em escala.”"],
        ["Model contracts", "“Apliquei contratos de dados nos marts — o build quebra se o schema mudar sem querer.”"],
        ["Unit tests (dbt 1.8)", "“Cobri a lógica de auditoria com unit tests e inputs mockados, além dos testes de dados.”"],
        ["Observabilidade", "“Tenho observabilidade: materializo falhas de teste e gravo metadados de execução.”"],
        ["CI + SQLFluff", "“PR roda dbt build + lint; em projeto grande, Slim CI com state:modified e deferral.”"],
        ["Masking de PII", "“Mascaro PII (placa) já na camada silver, por LGPD e menor exposição.”"],
        ["Exposures", "“Declaro o consumidor downstream (relatório de auditoria) no lineage.”"],
        ["Medallion + dbt", "“Estruturei o pipeline em camadas Medallion com dbt: bronze, silver e gold, com testes em cada camada.”"],
        ["Lookback (late-arriving)", "“Meu fato incremental tem janela de lookback com merge para não perder transações que chegam atrasadas — e documentei o trade-off.”"],
        ["Guarda contra fan-out", "“Meu join temporal poderia gerar fan-out se as vigências se sobrepusessem, então tranquei isso com um teste de ranges mutuamente exclusivos.”"],
        ["Semantic Layer (MetricFlow)", "“Defini métricas no Semantic Layer do dbt (revenue, suspect_rate, vazamento de receita) — uma definição governada, consultável por qualquer dimensão.”"],
        ["Sources + freshness", "“Declarei o contrato de ingestão com sources e source freshness (SLA de atraso), mostrando que penso o pipeline desde a entrada do dado, não só a transformação.”"],
        ["Observabilidade (Elementary)", "“Integrei o Elementary para observabilidade de dados — detecção de anomalia e tabelas de monitoramento no warehouse — e conheço os trade-offs de rodá-lo em DuckDB vs warehouse cloud.”"],
        ["Dinheiro em centavos inteiros", "“Agrego dinheiro em centavos inteiros e só converto para reais na exibição, pra não acumular erro de arredondamento/float — disciplina de dado financeiro.”"],
        ["Observabilidade agendada", "“Separei a detecção de anomalia num job agendado (tag observability), fora do CI de PR — monitoramento contínuo precisa de histórico, não é gate de merge.”"],
        ["Reconhecer o próprio risco", "“Documentei as limitações que eu mesmo encontrei (late-arriving, sources/freshness, observabilidade) num roadmap — antes de alguém perguntar.”"],
    ],
    widths=[2.3, 4.5],
)
para("Sugestão de frase-resumo para o topo do perfil:", bold=True)
para("“Engenheiro de Dados — construo pipelines de transformação com dbt (camadas "
     "Medallion, testes, contratos de dados, CI) portáveis para Databricks. Projeto de "
     "auditoria de pedágio com point-in-time, modelo incremental e unit tests.”",
     italic=True, color=AZUL)

# ---- rodapé final ----
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Fim do documento — dbt-toll-analytics. Pipeline validado com dbt build verde.")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = CINZA

out = os.path.join(os.path.dirname(__file__), "Documentacao_dbt_toll_analytics.docx")
doc.save(out)
print("OK:", out)
